import io
import re
import mistune
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from src.utils import logger


def clean_ai_metadata(content: str) -> str:
    """清理AI生成的元数据内容（如字数统计）

    与前端过滤逻辑保持一致：
    web/src/views/project/DocumentGenerator.vue:411-412
    """
    if not content:
        return content

    # 移除末尾的分隔符和字数统计块
    # 匹配从 --- 开始，到下一个标题或文档结尾的统计内容
    content = re.sub(r'---\s*\n\s*📊\s*字数统计：[\s\S]*?(?=\n#|\n\n|$)', '', content)

    # 移除独立的字数统计块（带emoji）
    # 匹配从 📊 字数统计：开始，到下一个标题或文档结尾的所有内容
    content = re.sub(r'[#\s-]*📊\s*字数统计：[\s\S]*?(?=\n#|\n\n|$)', '', content)

    # 移除独立的字数统计块（不带emoji，以# 开头）
    content = re.sub(r'\n#+\s*字数统计\s*\n[\s\S]*?(?=\n#|\n\n|$)', '', content)

    return content.strip()

class MarkdownToDocx:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """配置文档默认样式，特别是中文字体支持"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # 设置中文字体
        r = style.element.get_or_add_rPr()
        r.set(qn('w:eastAsia'), 'SimSun') # 宋体

    def convert(self, markdown_text, title=None):
        """将 Markdown 转换为 Docx 文件流"""
        try:
            # 清理AI元数据（双重保险）
            markdown_text = clean_ai_metadata(markdown_text)

            if title:
                # 添加主标题
                h = self.doc.add_heading(title, level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 使用 mistune 解析 Markdown 为 AST
            # mistune 3.x 默认支持插件
            markdown = mistune.create_markdown(renderer='ast')
            ast = markdown(markdown_text)

            for node in ast:
                self._process_node(node)

            # 保存到内存流
            file_stream = io.BytesIO()
            self.doc.save(file_stream)
            file_stream.seek(0)
            return file_stream
        except Exception as e:
            logger.error(f"Markdown 转 Docx 失败: {str(e)}")
            raise

    def _process_node(self, node):
        node_type = node.get('type')

        if node_type == 'heading':
            level = node.get('attrs', {}).get('level', 1)
            # 对于 level 0 已经用了标题的情况，调整后续级别
            p = self.doc.add_heading('', level=level)
            for child in node.get('children', []):
                self._process_inline_node(child, p)

        elif node_type == 'paragraph':
            p = self.doc.add_paragraph()
            for child in node.get('children', []):
                self._process_inline_node(child, p)

        elif node_type == 'list':
            ordered = node.get('attrs', {}).get('ordered', False)
            for item in node.get('children', []):
                self._process_list_item(item, ordered)

        elif node_type == 'block_code':
            # 代码块使用灰色背景和等宽字体
            p = self.doc.add_paragraph()
            p.style = 'No Spacing'
            run = p.add_run(node.get('raw', node.get('text', '')))
            run.font.name = 'Courier New'
            r = run._element.get_or_add_rPr()
            r.set(qn('w:eastAsia'), 'Courier New')

        elif node_type == 'thematic_break':
            self.doc.add_page_break()
            
        elif node_type == 'block_quote':
            for child in node.get('children', []):
                # 引用块在 docx 中较难完美还原，简单处理为缩进段落
                if child.get('type') == 'paragraph':
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(20)
                    for inline in child.get('children', []):
                        self._process_inline_node(inline, p)

    def _process_list_item(self, node, ordered):
        # 列表项通常包含段落
        style = 'List Number' if ordered else 'List Bullet'
        for child in node.get('children', []):
            if child.get('type') == 'paragraph':
                p = self.doc.add_paragraph(style=style)
                for inline in child.get('children', []):
                    self._process_inline_node(inline, p)
            else:
                # 递归处理可能的嵌套列表等
                self._process_node(child)

    def _process_inline_node(self, node, paragraph, is_bold=False, is_italic=False):
        node_type = node.get('type')
        
        if node_type == 'text':
            run = paragraph.add_run(node.get('raw', node.get('text', '')))
            if is_bold: run.bold = True
            if is_italic: run.italic = True
        
        elif node_type == 'strong':
            for child in node.get('children', []):
                self._process_inline_node(child, paragraph, is_bold=True, is_italic=is_italic)
        
        elif node_type == 'emphasis':
            for child in node.get('children', []):
                self._process_inline_node(child, paragraph, is_bold=is_bold, is_italic=True)
        
        elif node_type == 'codespan':
            run = paragraph.add_run(node.get('raw', node.get('text', '')))
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(199, 37, 78) # 类似 bootstrap 的代码颜色
            if is_bold: run.bold = True
            if is_italic: run.italic = True
        
        elif node_type == 'link':
            for child in node.get('children', []):
                self._process_inline_node(child, paragraph, is_bold=is_bold, is_italic=is_italic)
        
        elif node_type == 'linebreak':
            paragraph.add_run('\n')

        elif node_type == 'softbreak':
            paragraph.add_run(' ')

def export_markdown_to_docx(markdown_text, title=None):
    """便捷调用接口"""
    # 清理AI元数据
    cleaned_text = clean_ai_metadata(markdown_text)

    converter = MarkdownToDocx()
    return converter.convert(cleaned_text, title)
